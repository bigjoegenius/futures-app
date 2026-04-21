#!/usr/bin/env python3
"""
build_strategy_report.py — Generate the Word doc report from backtest_results in futures.db.

Usage:
  python build_strategy_report.py                # reads latest run_id, writes to Desktop
  python build_strategy_report.py --open         # also opens the file in Word when done
  python build_strategy_report.py --out ./report.docx

Sections:
  1. Executive summary (top 15 winning combos, bottom 5 losers)
  2. Fees / contract specs / funding note table
  3. Per-strategy deep dive (rules, winning markets/TFs, equity chart, notable losses)
  4. Per-market recommendations
  5. News that hurt us (top 20 losing trades with news tags grouped by event)
"""

from __future__ import annotations

import argparse
import os
import sqlite3
import subprocess
from datetime import datetime
from collections import defaultdict

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from futures_config import DB_PATH, FUTURES
from market_analyzer import STRATEGIES
from strategy_engine import CONTRACT_SPECS, FEE_PER_CONTRACT_PER_SIDE, SLIPPAGE_TICKS


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
BT_DIR = os.path.join(BASE_DIR, "backtests")


# ─── helpers ────────────────────────────────────────────────────────────
def _shade_cell(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    return h


def _latest_run_id(conn) -> str | None:
    try:
        row = conn.execute("SELECT run_id FROM backtest_results ORDER BY completed_at DESC LIMIT 1").fetchone()
        return row[0] if row else None
    except sqlite3.OperationalError:
        return None


def _load_results_from_json() -> list[dict]:
    """Fallback loader: read per-combo JSON files from backtests/ dir."""
    import glob, json as _json
    results = []
    for path in glob.glob(os.path.join(BT_DIR, "*.json")):
        if os.path.basename(path) == "run_summary.json":
            continue
        try:
            with open(path) as f:
                data = _json.load(f)
            m = data.get("metrics", {})
            results.append({
                "strategy": data.get("strategy"),
                "symbol": data.get("symbol"),
                "timeframe": data.get("timeframe"),
                "bars": data.get("bars", 0),
                "trades": m.get("trades", 0),
                "wins": m.get("wins", 0),
                "losses": m.get("losses", 0),
                "win_rate": m.get("win_rate", 0),
                "total_pnl": m.get("total_pnl", 0),
                "final_balance": m.get("final_balance", 10_000),
                "profit_factor": m.get("profit_factor", 0),
                "max_drawdown": m.get("max_drawdown", 0),
                "sharpe": m.get("sharpe", 0),
                "avg_win": m.get("avg_win", 0),
                "avg_loss": m.get("avg_loss", 0),
                "trades_detail": data.get("trades", []),
            })
        except Exception:
            continue
    results.sort(key=lambda r: r["total_pnl"], reverse=True)
    return results


def _load_results(conn, run_id: str) -> list[dict]:
    try:
        cur = conn.execute(
            "SELECT strategy, symbol, timeframe, bars, trades, wins, losses, win_rate, "
            "total_pnl, final_balance, profit_factor, max_drawdown, sharpe, avg_win, avg_loss "
            "FROM backtest_results WHERE run_id=? ORDER BY total_pnl DESC",
            (run_id,),
        )
        cols = ["strategy","symbol","timeframe","bars","trades","wins","losses","win_rate",
                "total_pnl","final_balance","profit_factor","max_drawdown","sharpe","avg_win","avg_loss"]
        return [dict(zip(cols, r)) for r in cur.fetchall()]
    except sqlite3.OperationalError:
        return _load_results_from_json()


def _load_trades_for(conn, run_id: str, strategy: str, symbol: str, timeframe: str) -> list[dict]:
    try:
        cur = conn.execute(
            "SELECT direction, entry_time, entry_price, exit_time, exit_price, pnl_dollars, "
            "pnl_pct, fees, exit_reason, news_tags "
            "FROM backtest_trades "
            "WHERE run_id=? AND strategy=? AND symbol=? AND timeframe=? "
            "ORDER BY entry_time",
            (run_id, strategy, symbol, timeframe),
        )
        cols = ["direction","entry_time","entry_price","exit_time","exit_price","pnl_dollars",
                "pnl_pct","fees","exit_reason","news_tags"]
        return [dict(zip(cols, r)) for r in cur.fetchall()]
    except sqlite3.OperationalError:
        # Fallback to JSON
        import json as _json
        combo_key = f"{strategy}_{symbol.replace('=F','')}_{timeframe}"
        path = os.path.join(BT_DIR, f"{combo_key}.json")
        if not os.path.exists(path):
            return []
        try:
            with open(path) as f:
                return _json.load(f).get("trades", [])
        except Exception:
            return []


def _load_losing_trades_with_news(conn, run_id: str, limit: int = 60) -> list[dict]:
    try:
        cur = conn.execute(
            "SELECT strategy, symbol, timeframe, direction, entry_time, exit_time, "
            "entry_price, exit_price, pnl_dollars, news_tags "
            "FROM backtest_trades WHERE run_id=? AND pnl_dollars < 0 AND news_tags != '' "
            "ORDER BY pnl_dollars ASC LIMIT ?",
            (run_id, limit),
        )
        cols = ["strategy","symbol","timeframe","direction","entry_time","exit_time",
                "entry_price","exit_price","pnl_dollars","news_tags"]
        return [dict(zip(cols, r)) for r in cur.fetchall()]
    except sqlite3.OperationalError:
        # Fallback: scan JSON files
        import glob, json as _json
        all_trades = []
        for path in glob.glob(os.path.join(BT_DIR, "*.json")):
            try:
                with open(path) as f:
                    d = _json.load(f)
                for t in d.get("trades", []):
                    if t.get("pnl_dollars", 0) < 0 and t.get("news_tags"):
                        t2 = dict(t)
                        t2["strategy"] = d.get("strategy")
                        t2["symbol"]   = d.get("symbol")
                        t2["timeframe"]= d.get("timeframe")
                        all_trades.append(t2)
            except Exception:
                continue
        all_trades.sort(key=lambda t: t.get("pnl_dollars", 0))
        return all_trades[:limit]


# ─── section builders ──────────────────────────────────────────────────
def section_cover(doc, run_id: str, results: list[dict]):
    wins = [r for r in results if r["total_pnl"] > 0]
    losses = [r for r in results if r["total_pnl"] <= 0]
    total_trades = sum(r["trades"] for r in results)

    title = doc.add_heading("Futures Strategy Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Backtest run {run_id} — generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    r.italic = True
    r.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("28 strategies · 18 markets · 5 timeframes\n").bold = True
    p.add_run(f"{len(results)} combinations backtested — {len(wins)} profitable / {len(losses)} unprofitable · "
              f"{total_trades:,} simulated trades total\n").font.size = Pt(11)
    p.add_run("Starting balance per combo: $10,000  ·  Risk mode: moderate (1% per trade)  ·  "
              "Fees: $2.50/contract/side + 1-tick slippage")


def section_summary(doc, results: list[dict]):
    _add_heading(doc, "Executive Summary", level=1)

    doc.add_paragraph(
        "The 28-strategy catalog is documented in strategies/STRATEGY_CATALOG.md. "
        "Every market has at least one specialist strategy plus the 10 market-agnostic ones. "
        "Every timeframe (1m, 5m, 15m, 1h, 1d) has multiple strategies assigned. "
        "Results below are walk-forward: each bar's decision uses only prior-bar data."
    )

    top = sorted(results, key=lambda r: r["total_pnl"], reverse=True)[:15]
    _add_heading(doc, "Top 15 winning combinations", level=2)
    tbl = doc.add_table(rows=1, cols=8)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Strategy","Market","TF","Trades","WR %","P&L $","PF","Sharpe"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for r in top:
        row = tbl.add_row().cells
        row[0].text = r["strategy"]
        row[1].text = r["symbol"]
        row[2].text = r["timeframe"]
        row[3].text = f"{r['trades']}"
        row[4].text = f"{r['win_rate']:.1f}"
        row[5].text = f"{r['total_pnl']:+,.0f}"
        row[6].text = f"{r['profit_factor']:.2f}"
        row[7].text = f"{r['sharpe']:.2f}"
        _shade_cell(row[5], "DCFFE4")

    doc.add_paragraph()
    worst = sorted([r for r in results if r["trades"] > 0], key=lambda r: r["total_pnl"])[:10]
    _add_heading(doc, "Worst 10 combinations (avoid)", level=2)
    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Light Grid Accent 2"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Strategy","Market","TF","Trades","WR %","P&L $","Max DD %"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for r in worst:
        row = tbl.add_row().cells
        row[0].text = r["strategy"]
        row[1].text = r["symbol"]
        row[2].text = r["timeframe"]
        row[3].text = f"{r['trades']}"
        row[4].text = f"{r['win_rate']:.1f}"
        row[5].text = f"{r['total_pnl']:+,.0f}"
        row[6].text = f"{r['max_drawdown']:.1f}"
        _shade_cell(row[5], "FFDCE0")


def section_fees(doc):
    _add_heading(doc, "Fees, Slippage & Contract Specs", level=1)
    doc.add_paragraph(
        f"Assumed per-side commission: ${FEE_PER_CONTRACT_PER_SIDE:.2f}/contract. "
        f"Slippage: {SLIPPAGE_TICKS} tick against fill on entry and exit. "
        "These are approximations for CME futures via retail brokers (Apex, NinjaTrader, TopStep, etc.)."
    )
    doc.add_paragraph().add_run(
        "Futures vs crypto — important distinction: futures have NO funding rate. "
        "Instead, the cost of carry shows up as (a) roll costs when front-month expiration nears "
        "(approximately 1 tick per calendar roll, modeled as a small trade cost), and "
        "(b) contango/backwardation embedded in the term structure itself. "
        "For the backtests above we use the continuous front-month price series from yfinance, "
        "so roll effects are already baked in."
    ).italic = True

    _add_heading(doc, "Contract specifications", level=2)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light List Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Symbol","Product","Point value $","Tick","Tick value $"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for sym, sp in CONTRACT_SPECS.items():
        row = tbl.add_row().cells
        row[0].text = sym
        row[1].text = sp["name"]
        row[2].text = f"{sp['point_value']:,.2f}"
        row[3].text = f"{sp['tick']:g}"
        row[4].text = f"{sp['point_value'] * sp['tick']:.2f}"


def section_per_strategy(doc, results: list[dict], conn, run_id: str):
    _add_heading(doc, "Per-Strategy Deep Dive", level=1)
    doc.add_paragraph("One entry per strategy with its rules, best + worst (market, timeframe) combinations, and news-context for its biggest losers.")
    by_strategy = defaultdict(list)
    for r in results:
        by_strategy[r["strategy"]].append(r)

    for sid, spec in STRATEGIES.items():
        doc.add_page_break()
        _add_heading(doc, f"{sid} — {spec.get('name','')}", level=2)
        p = doc.add_paragraph()
        p.add_run("Rules: ").bold = True
        p.add_run(spec.get("description", ""))
        p.add_run(f"\nDirection: {spec.get('direction','BOTH')}  |  "
                  f"Timeframes: {', '.join(spec.get('timeframes', []))}  |  "
                  f"Markets: {', '.join(spec.get('markets', ['all']))}"
                  f"\nStop mult: {spec.get('stop_atr_mult',1.2)}× ATR · "
                  f"Target mult: {spec.get('target_atr_mult',1.8)}× ATR")

        combos = sorted(by_strategy.get(sid, []), key=lambda r: r["total_pnl"], reverse=True)
        if not combos:
            doc.add_paragraph("No backtest data yet.")
            continue

        # Combo summary table
        tbl = doc.add_table(rows=1, cols=7)
        tbl.style = "Light Grid Accent 3"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Market","TF","Trades","WR %","P&L $","PF","Sharpe"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for r in combos[:18]:
            row = tbl.add_row().cells
            row[0].text = r["symbol"]
            row[1].text = r["timeframe"]
            row[2].text = f"{r['trades']}"
            row[3].text = f"{r['win_rate']:.1f}"
            row[4].text = f"{r['total_pnl']:+,.0f}"
            row[5].text = f"{r['profit_factor']:.2f}"
            row[6].text = f"{r['sharpe']:.2f}"
            _shade_cell(row[4], "DCFFE4" if r["total_pnl"] > 0 else "FFDCE0")

        # Best equity chart
        best = combos[0] if combos[0]["total_pnl"] > 0 else None
        if best:
            combo_key = f"{best['strategy']}_{best['symbol'].replace('=F','')}_{best['timeframe']}"
            png_path = os.path.join(BT_DIR, f"equity_{combo_key}.png")
            if os.path.exists(png_path):
                doc.add_paragraph()
                doc.add_picture(png_path, width=Inches(6.0))
                p = doc.add_paragraph(f"Best combo equity curve: {best['symbol']} {best['timeframe']}")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.italic = True
                    r.font.size = Pt(9)

        # Losing trades with news tags (for the strategy's worst combo)
        worst = [c for c in combos if c["total_pnl"] < 0]
        if worst:
            w = worst[0]
            trades = _load_trades_for(conn, run_id, sid, w["symbol"], w["timeframe"])
            notable = [t for t in trades if t["news_tags"] and t["pnl_dollars"] < 0][:5]
            if notable:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run(f"Losing trades on {w['symbol']} {w['timeframe']} that overlapped news events:").bold = True
                for t in notable:
                    doc.add_paragraph(
                        f"  · {t['direction']} {t['entry_time'][:16]} → {t['exit_time'][:16] if t['exit_time'] else '--'} "
                        f"P&L ${t['pnl_dollars']:+.0f} · {t['news_tags']}",
                        style="List Bullet",
                    )


def section_per_market(doc, results: list[dict]):
    doc.add_page_break()
    _add_heading(doc, "Per-Market Recommendations", level=1)
    doc.add_paragraph("For each market, the recommended strategy set is ranked by Sharpe (risk-adjusted). Deploy the top 2-3 in the autopilot's default set.")
    by_market = defaultdict(list)
    for r in results:
        by_market[r["symbol"]].append(r)
    for sym in FUTURES:
        combos = sorted(by_market.get(sym, []), key=lambda r: r.get("sharpe", 0), reverse=True)
        if not combos:
            continue
        _add_heading(doc, f"{sym} — {FUTURES[sym]}", level=2)
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Light List Accent 2"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Strategy","TF","Trades","WR %","P&L $","Sharpe"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for r in combos[:10]:
            row = tbl.add_row().cells
            row[0].text = r["strategy"]
            row[1].text = r["timeframe"]
            row[2].text = f"{r['trades']}"
            row[3].text = f"{r['win_rate']:.1f}"
            row[4].text = f"{r['total_pnl']:+,.0f}"
            row[5].text = f"{r['sharpe']:.2f}"
            _shade_cell(row[4], "DCFFE4" if r["total_pnl"] > 0 else ("FFF4D6" if r["total_pnl"] == 0 else "FFDCE0"))


def section_news(doc, losing_trades: list[dict]):
    doc.add_page_break()
    _add_heading(doc, "News That Hurt Us", level=1)
    doc.add_paragraph(
        "Losing trades whose entry or exit window (±4 hours) overlapped a known high-impact event — "
        "FOMC decisions, WASDE reports, OPEC meetings, or major geopolitical shocks. "
        "Strategies should consider blocking entries in the ±1 day window around these events in the future."
    )
    if not losing_trades:
        doc.add_paragraph("No losing trades overlapped a known news event in this run.")
        return

    # Group by first tag type
    by_event = defaultdict(list)
    for t in losing_trades:
        first_tag = (t["news_tags"] or "").split(",")[0]
        head = first_tag.split("(")[0] or "Other"
        by_event[head].append(t)

    for head, items in sorted(by_event.items(), key=lambda kv: -len(kv[1])):
        _add_heading(doc, f"{head} — {len(items)} losing trades", level=2)
        tbl = doc.add_table(rows=1, cols=7)
        tbl.style = "Light Grid Accent 4"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Strategy","Sym","TF","Dir","Entry","Exit","P&L $"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for t in items[:12]:
            row = tbl.add_row().cells
            row[0].text = t["strategy"]
            row[1].text = t["symbol"]
            row[2].text = t["timeframe"]
            row[3].text = t["direction"]
            row[4].text = t["entry_time"][:16] if t["entry_time"] else "--"
            row[5].text = t["exit_time"][:16] if t["exit_time"] else "--"
            row[6].text = f"{t['pnl_dollars']:+.0f}"
            _shade_cell(row[6], "FFDCE0")


# ─── main ──────────────────────────────────────────────────────────────
def build(out_path: str, open_after: bool = False):
    conn = sqlite3.connect(DB_PATH)
    run_id = _latest_run_id(conn) or "json_fallback"
    results = _load_results(conn, run_id)
    if not results:
        # Full JSON fallback
        results = _load_results_from_json()
    if not results:
        print("No backtest results yet (neither DB nor JSON files). Run `python backtest_all.py` first.")
        return
    losing_trades = _load_losing_trades_with_news(conn, run_id)
    conn.close()

    doc = Document()
    # Page margins
    for s in doc.sections:
        s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)
        s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)

    section_cover(doc, run_id, results)
    section_summary(doc, results)
    section_fees(doc)
    conn = sqlite3.connect(DB_PATH)
    section_per_strategy(doc, results, conn, run_id)
    conn.close()
    section_per_market(doc, results)
    section_news(doc, losing_trades)

    doc.save(out_path)
    print(f"Saved: {out_path}  ({os.path.getsize(out_path)/1024:.0f} KB)")
    if open_after:
        try:
            subprocess.run(["open", out_path])
        except Exception:
            pass


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=None)
    ap.add_argument("--open", action="store_true")
    args = ap.parse_args()

    if args.out:
        out = args.out
    else:
        desk = os.path.expanduser("~/Desktop")
        stamp = datetime.now().strftime("%Y-%m-%d")
        out = os.path.join(desk, f"Futures_Strategy_Report_{stamp}.docx")
    build(out, open_after=args.open)


if __name__ == "__main__":
    main()
